#
# Copyright 2021 Open Raven Inc. and the Mockingbird project authors
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#


from typing import final

from docx import Document

from mockingbird.__base import __BaseDocument
from .__base import __BaseUnstructuredDataType


class DOCXDocument(__BaseDocument):
    def __init__(self):
        super().__init__(extension="docx")

        # Create a list of docx formats we're going to export.
        self._docx_styles = []
        active_styles = self._configurable_dict["unstructured_data"]["docx_document"]["active_styles"]

        if active_styles["paragraph_style"]:
            self._docx_styles.append(_DocxParagraphStyle)

        if active_styles["footer_style"]:
            self._docx_styles.append(_DocxFooterStyle)

        if active_styles["bullet_point_style"]:
            self._docx_styles.append(_DocxBulletPointStyle)

        if active_styles["chat_style"]:
            self._docx_styles.append(_DocxChatStyle)

    @final
    def save(self, save_path: str) -> None:

        for style in self._docx_styles:
            instantiated_style = style()
            instantiated_style.clone_sensitive_data(other=self)
            instantiated_style.save(save_path=save_path)
            self._meta_data_object.add_other_meta_data(instantiated_style._meta_data_object)


class _DocxParagraphStyle(__BaseUnstructuredDataType):
    """
    Writes a simple paragraph containing sensitive-soup.
    """

    def __init__(self):
        super().__init__(extension="docx")

    @final
    def save(self, save_path: str) -> None:
        """

        """
        save_file = self.setup_save_file(save_path=save_path, extension=self.extension)

        document = Document()
        document.add_heading('Paragraph Styled Document', 0)

        sensitive_soup = self._get_sensitive_soup()
        document.add_paragraph(sensitive_soup)
        document.save(save_file)
        self._log_save(save_file)


class _DocxFooterStyle(__BaseUnstructuredDataType):
    """
    Writes a simple document with sensitive-soup in the footer.
    """

    def __init__(self):
        super().__init__(extension="docx")

    @final
    def save(self, save_path: str) -> None:
        """

        """
        save_file = self.setup_save_file(save_path=save_path, extension=self.extension)
        sensitive_soup = self._get_sensitive_soup()

        document = Document()
        document.add_heading('Sensitive-Data in Footer Styled Document', 0)

        section = document.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = sensitive_soup

        document.save(save_file)
        self._log_save(save_file)


class _DocxBulletPointStyle(__BaseUnstructuredDataType):
    """
    Writes a simple document with sensitive-soup in the footer.
    """

    def __init__(self):
        super().__init__(extension="docx")

    @final
    def save(self, save_path: str) -> None:
        """

        """
        save_file = self.setup_save_file(save_path=save_path, extension=self.extension)
        enumerated_groups = self._get_enumerated_style()

        document = Document()
        document.add_heading('Sensitive Data Stored in Bullet Points', 0)

        for group in enumerated_groups:
            key, enumerated_items = group

            document.add_heading(key, level=1)

            for item in enumerated_items:
                document.add_paragraph(item, style="List Bullet")

        document.save(save_file)
        self._log_save(save_file)


class _DocxChatStyle(__BaseUnstructuredDataType):
    """
    Writes a simple document with sensitive-soup in the footer.
    """

    def __init__(self):
        super().__init__(extension="docx")

    @final
    def save(self, save_path: str) -> None:
        """

        """
        save_file = self.setup_save_file(save_path=save_path, extension=self.extension)
        chat_log = self._get_chat_log()

        document = Document()
        document.add_heading('A chat between two people', 0)

        for line in chat_log:
            document.add_paragraph(line)

        document.save(save_file)
        self._log_save(save_file)
